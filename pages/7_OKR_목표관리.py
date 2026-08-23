import os
import io
import re
import json
import hashlib
import uuid
import mimetypes
from datetime import date, datetime, timedelta

import pandas as pd
import requests
import streamlit as st
from supabase import create_client

st.set_page_config(page_title="OKR 목표관리", layout="wide")

# ── 비밀번호 게이트 ───────────────────────────────────────────
PW = os.environ.get("APP_PASSWORD")
if PW and not st.session_state.get("ok"):
    pw = st.text_input("비밀번호", type="password")
    if st.button("입장"):
        if pw == PW:
            st.session_state.ok = True; st.rerun()
        else:
            st.error("비밀번호가 올바르지 않습니다.")
    st.stop()

@st.cache_resource
def sb():
    url = os.environ.get("SUPABASE_URL"); key = os.environ.get("SUPABASE_SERVICE_KEY")
    if not url or not key:
        st.error("❌ SUPABASE_URL / SUPABASE_SERVICE_KEY 환경변수가 없습니다."); st.stop()
    return create_client(url, key)
SUPA = sb()

ADMIN_EMAIL = "jhw@slam-global.com"

ATTACH_BUCKET = "okr-attachments"

# ── 직무내역서 PDF용 한글 폰트 (repo 루트/assets 기준 절대경로) ──
_JD_BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_JD_FONT_REGULAR = os.path.join(_JD_BASE_DIR, "assets", "NanumGothic-Regular.ttf")
_JD_FONT_BOLD = os.path.join(_JD_BASE_DIR, "assets", "NanumGothic-Bold.ttf")

def upload_photo(file, person):
    """이미지를 Supabase Storage에 올리고 공개 URL을 반환한다."""
    ext = (file.name.rsplit(".", 1)[-1] if "." in file.name else "jpg").lower()
    path = f"{person}/{uuid.uuid4().hex}.{ext}"
    data = file.read()
    content_type = mimetypes.guess_type(file.name)[0] or "application/octet-stream"
    SUPA.storage.from_(ATTACH_BUCKET).upload(path, data, {"content-type": content_type})
    return SUPA.storage.from_(ATTACH_BUCKET).get_public_url(path)

def render_attachments_editor(item, key_prefix):
    """업무 항목 하나에 링크/사진을 추가하고, 이미 첨부된 것들을 보여주는 위젯."""
    atts = item.get("attachments") or []
    if atts:
        for idx, a in enumerate(atts):
            with st.container(border=True):
                ac1, ac2 = st.columns([5, 1])
                if a.get("type") == "photo":
                    ac1.image(a["url"], width=180)
                    ac1.caption(a.get("label") or "완료 스크린샷")
                else:
                    ac1.markdown(f"🔗 [{a.get('label') or a['url']}]({a['url']})")
                if ac2.button("삭제", key=f"delatt_{key_prefix}_{idx}"):
                    new_atts = [x for j, x in enumerate(atts) if j != idx]
                    SUPA.table("okr_items").update({"attachments": new_atts}).eq("id", item["id"]).execute()
                    refresh()
    with st.form(f"add_attach_{key_prefix}", clear_on_submit=True):
        st.caption("🔗 링크 또는 📷 완료 화면 캡처 추가")
        a1, a2 = st.columns(2)
        link_url = a1.text_input("링크 URL (선택)")
        link_label = a2.text_input("링크 설명 (선택)")
        photo_file = st.file_uploader("사진 첨부 (선택)", type=["png", "jpg", "jpeg", "webp"], key=f"photofile_{key_prefix}")
        if st.form_submit_button("➕ 첨부 추가"):
            new_atts = list(atts)
            added = False
            if link_url.strip():
                new_atts.append({"type": "link", "url": link_url.strip(), "label": link_label.strip(), "added_at": datetime.utcnow().isoformat()})
                added = True
            if photo_file is not None:
                try:
                    url = upload_photo(photo_file, item["person"])
                    new_atts.append({"type": "photo", "url": url, "label": photo_file.name, "added_at": datetime.utcnow().isoformat()})
                    added = True
                except Exception as e:
                    st.error(f"사진 업로드 실패: {e}")
            if added:
                SUPA.table("okr_items").update({"attachments": new_atts}).eq("id", item["id"]).execute()
                st.success("첨부를 추가했습니다."); refresh()
            else:
                st.warning("링크나 사진을 하나 이상 입력해주세요.")

def send_via_resend(to_addr, subject, body_text, purpose="manual"):
    """Resend API로 이메일 발송 + email_log 테이블에 기록. (RESEND_API_KEY 환경변수 필요)"""
    api_key = os.environ.get("RESEND_API_KEY")
    sender = os.environ.get("RESEND_FROM", "브랜드슬램 OKR <onboarding@resend.dev>")
    if not api_key:
        return False, "RESEND_API_KEY 환경변수가 설정되어 있지 않습니다."
    try:
        res = requests.post(
            "https://api.resend.com/emails",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={"from": sender, "to": [to_addr], "subject": subject, "text": body_text},
            timeout=15,
        )
        ok = res.status_code < 300
        err = None if ok else f"{res.status_code} {res.text}"[:500]
    except Exception as e:
        ok, err = False, str(e)
    try:
        SUPA.table("email_log").insert({
            "purpose": purpose, "recipient": to_addr, "subject": subject, "body": body_text,
            "status": "sent" if ok else "failed", "error": err,
        }).execute()
    except Exception:
        pass
    return ok, err


def build_jd_pdf(position_title, org_tag, objective, categorized_items):
    """담당자 1명의 직무내역서를 PDF bytes로 만든다.
    categorized_items: {카테고리: [(업무명, 주기라벨), ...]} — 심플하게 이 3가지 정보만 담는다
    (수치·진행률·급여 등 민감/내부 정보는 절대 포함하지 않음 — 채용 후보자에게 그대로 전달되는 문서라서).
    reportlab의 기본 CID 폰트 대신, repo에 포함된 TTF를 직접 등록해서 사용한다
    (Railway 환경에서 CID 폰트 매핑 오류가 났던 과거 이슈를 우회하기 위함)."""
    import io as _io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors as rl_colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import ParagraphStyle

    if "Nanum" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("Nanum", _JD_FONT_REGULAR))
        pdfmetrics.registerFont(TTFont("Nanum-Bold", _JD_FONT_BOLD))

    style_title = ParagraphStyle("jd_title", fontName="Nanum-Bold", fontSize=20, leading=26, textColor=rl_colors.HexColor("#16181d"))
    style_sub = ParagraphStyle("jd_sub", fontName="Nanum", fontSize=10.5, leading=15, textColor=rl_colors.HexColor("#6b7280"))
    style_h2 = ParagraphStyle("jd_h2", fontName="Nanum-Bold", fontSize=13, leading=18, textColor=rl_colors.HexColor("#16181d"), spaceBefore=14, spaceAfter=6)
    style_body = ParagraphStyle("jd_body", fontName="Nanum", fontSize=10.5, leading=16, textColor=rl_colors.HexColor("#2b2f36"))
    style_cat = ParagraphStyle("jd_cat", fontName="Nanum-Bold", fontSize=11, leading=15, textColor=rl_colors.HexColor("#16181d"), spaceBefore=10, spaceAfter=4)
    style_foot = ParagraphStyle("jd_foot", fontName="Nanum", fontSize=8.5, leading=12, textColor=rl_colors.HexColor("#9aa0a8"))

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, topMargin=22 * mm, bottomMargin=18 * mm, leftMargin=20 * mm, rightMargin=20 * mm,
        title=f"{position_title} 직무내역서",
    )

    header_line = "브랜드슬램 (Brandslam)" + (f" · {org_tag}" if org_tag else "") + f" · {position_title}"
    story = [
        Paragraph("직무내역서", style_title),
        Paragraph(header_line, style_sub),
        Paragraph(f"발행일 {date.today().isoformat()}", style_sub),
        Spacer(1, 8 * mm),
    ]
    if objective:
        story.append(Paragraph("직무개요", style_h2))
        story.append(Paragraph(objective, style_body))

    story.append(Paragraph("주요 업무", style_h2))
    if not categorized_items:
        story.append(Paragraph("포함된 업무가 없습니다.", style_body))
    else:
        for cat, tasks in categorized_items.items():
            story.append(Paragraph(f"● {cat}", style_cat))
            tbl_data = [["업무명", "수행 주기"]] + [[t, c] for t, c in tasks]
            tbl = Table(tbl_data, colWidths=[110 * mm, 40 * mm])
            tbl.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, -1), "Nanum"),
                ("FONTNAME", (0, 0), (-1, 0), "Nanum-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#f2f3f5")),
                ("LINEBELOW", (0, 0), (-1, 0), 0.75, rl_colors.HexColor("#dfe2e8")),
                ("LINEBELOW", (0, 1), (-1, -1), 0.5, rl_colors.HexColor("#eef0f3")),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.append(tbl)

    story.append(Spacer(1, 14 * mm))
    story.append(Paragraph("본 문서는 브랜드슬램 내부 OKR 관리 시스템에서 자동 생성되었습니다.", style_foot))

    doc.build(story)
    return buf.getvalue()


# ── 리포트 업로드 → LLM 기반 KPI 자동 추출 ────────────────────
def extract_text_from_upload(uploaded_file):
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    if name.endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    try:
        return data.decode("utf-8")
    except Exception:
        return data.decode("utf-8", errors="ignore")

def _parse_json_array(text):
    m = re.search(r"\[.*\]", text, re.S)
    if not m:
        raise ValueError("응답에서 JSON 배열을 찾지 못했습니다: " + text[:200])
    return json.loads(m.group(0))

_KPI_UNIT_WORDS = ("건", "명", "회", "%", "억", "만원", "개", "차", "매출", "목표", "수량", "월", "주", "분기",
                    "달성", "이상", "이하", "협업", "진행", "구축", "확보", "운영", "관리", "리스트", "제공")
_BULLET_MARKERS = re.compile(r"^\s*([\*\-•▪◦]|\d+[\)\.]|\(?\d+\)|[①②③④⑤⑥⑦⑧⑨])")

def filter_numeric_lines(text):
    """토큰을 줄이되 '누락'이 훨씬 더 큰 문제이므로 회수율(recall)을 우선한다.
    숫자/단위 표현이 있는 줄, 목록·번호 형태의 줄, 짧은 줄(제목/헤더일 가능성)은 전부 남기고,
    확실히 긴 순수 서술형 문단(80자 이상이면서 위 어떤 조건도 안 걸리는 줄)만 잘라낸다."""
    kept = []
    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue
        has_digit_or_unit = bool(re.search(r"\d", s)) or any(w in s for w in _KPI_UNIT_WORDS)
        looks_like_list_item = bool(_BULLET_MARKERS.match(s))
        is_short = len(s) <= 80
        if has_digit_or_unit or looks_like_list_item or is_short:
            kept.append(s)
    return "\n".join(kept)

def content_hash_of(text):
    return hashlib.sha256(text.strip().encode("utf-8")).hexdigest()

def find_cached_extraction(hash_value):
    """완전히 같은 문서(hash 동일)를 예전에 이미 분석한 적 있으면 그 결과를 그대로 재사용 —
    API를 다시 호출하지 않아서 토큰을 아낀다."""
    if not hash_value:
        return None
    rows = SUPA.table("okr_report_uploads").select("extracted_json") \
        .eq("content_hash", hash_value).eq("applied", True) \
        .order("created_at", desc=True).limit(1).execute().data
    return rows[0]["extracted_json"] if rows else None

def call_claude_extract(person, raw_text, few_shot):
    """Anthropic API로 문서에서 정량 KPI 항목을 뽑아 JSON 리스트로 반환.
    few_shot: 그 사람의 과거 업로드 중 실제로 반영됐던 (원문, 추출결과) 몇 건 — 스타일 일관성을 위한 참고자료."""
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return None, "ANTHROPIC_API_KEY 환경변수가 설정되어 있지 않습니다."
    system = (
        "너는 한국 스타트업의 업무계획서/KPI 보고서에서 '측정 가능한 목표'를 최대한 빠짐없이 뽑아내는 도우미다. "
        "가장 중요한 원칙: **놓치는 것이 훨씬 나쁘다.** 확신이 안 서도 일단 후보로 포함시켜라 — "
        "사람이 나중에 검토 화면에서 체크 해제하면 되니, 너는 과감하게 넓게 뽑아라. "
        "숫자가 명시되지 않은 목표(예: '~와 협업 진행', 'DB 구축')도 target_qty를 0으로 넣어 반드시 포함해라. "
        "글머리 기호(*, -, 1), 숫자.)로 시작하는 줄, 소제목처럼 보이는 짧은 줄들도 놓치지 말고 전부 검토해라. "
        "출력은 오직 JSON 배열이어야 하고, 그 외 설명 텍스트는 절대 포함하지 않는다. "
        "각 항목 형식: {\"category\": 짧은 분류명, \"title\": 업무명(간결하게), "
        "\"target_qty\": 숫자(문서에 수치가 없으면 0), \"unit\": 단위(건/명/회 등, 수치 없으면 \"건\"), "
        "\"cadence\": \"weekly\"|\"monthly\"|\"quarterly\"|\"once\" 중 하나, "
        "\"note\": 원문에서 이 항목의 근거가 된 문구 한 줄}. "
        "월 N건처럼 명시된 것은 monthly, 주 N회는 weekly, 분기/누적 목표는 quarterly, "
        "일회성 프로젝트/마감이 있는 건은 once로 분류해라."
    )
    messages = []
    for ex in (few_shot or [])[:2]:
        ex_text = filter_numeric_lines(ex.get("raw_text") or "") or (ex.get("raw_text") or "")[:600]
        messages.append({"role": "user", "content": f"문서:\n{ex_text[:800]}"})
        messages.append({"role": "assistant", "content": json.dumps(ex.get("extracted_json") or [], ensure_ascii=False)})
    filtered = filter_numeric_lines(raw_text)
    send_text = filtered if len(filtered) >= 30 else raw_text[:6000]
    messages.append({"role": "user", "content": f"담당자: {person}\n\n문서 내용 (목표와 무관해 보이는 순수 서술 문단만 일부 제거됨, 나머지는 원문 그대로):\n{send_text[:8000]}"})
    try:
        res = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01", "content-type": "application/json"},
            json={"model": "claude-sonnet-5", "max_tokens": 4000, "system": system, "messages": messages},
            timeout=60,
        )
        if res.status_code >= 300:
            return None, f"{res.status_code} {res.text[:300]}"
        data = res.json()
        text = "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text")
        items = _parse_json_array(text)
        return items, None
    except Exception as e:
        return None, str(e)

# ── 촘촘한 시트형 UI (Notion/Linear/엑셀 느낌) ────────────────
st.markdown("""
<style>
/* 상세보기 expander를 '펼치는 시트 행'처럼 컴팩트하게 */
div[data-testid="stExpander"] {
    border: none !important;
    border-top: none !important;
    border-bottom: 1px solid #eef0f3 !important;
    border-radius: 0 !important;
    margin-bottom: 2px !important;
}
div[data-testid="stExpander"] summary {
    padding: 4px 6px !important;
    min-height: unset !important;
}
div[data-testid="stExpander"] summary:hover { background: #f7f8fa !important; }
div[data-testid="stExpander"] summary p { font-size: 12px !important; color: #9aa0a8 !important; }
div[data-testid="stExpander"] div[data-testid="stExpanderDetails"] { padding-top: 4px !important; }

/* 시트 행 헤더 */
.okr-hdr {
    display: grid; gap: 10px; padding: 6px 8px;
    font-size: 11px; font-weight: 700; color: #8a8f98;
    border-bottom: 1.5px solid #dfe2e8;
    text-transform: uppercase; letter-spacing: .03em;
}
/* 시트 행 본문 */
.okr-row-html {
    display: grid; gap: 10px; padding: 9px 8px;
    align-items: center; font-size: 13.5px;
    border-bottom: 1px solid #f2f3f5;
}
.okr-row-html:hover { background: #fafbfc; }
.okr-title { font-weight: 700; }
.okr-cat { color: #9aa0a8; font-size: 11px; margin-left: 7px; font-weight: 500; }
.okr-num { font-variant-numeric: tabular-nums; color: #444; }
.okr-person { font-weight: 700; font-size: 12.5px; }
</style>
""", unsafe_allow_html=True)

st.title("🎯 OKR 목표관리")
st.caption("OWM → LSP → 브랜드슬램 비전 캐스케이드와 구성원별 OKR · 실시간 DB 연동")

# ── 관리자 확인 (이메일 + 비밀번호) ──────────────────────────
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD")
st.session_state.setdefault("user_email", "")
st.session_state.setdefault("admin_pw", "")
id_c1, id_c2 = st.columns([3, 1.4])
with id_c2:
    st.session_state["user_email"] = st.text_input(
        "내 이메일", value=st.session_state["user_email"],
        placeholder="jhw@slam-global.com", help="관리자 이메일 + 관리자 비밀번호를 모두 맞게 입력해야 관리자 모드가 됩니다."
    )
    email_matches = st.session_state["user_email"].strip().lower() == ADMIN_EMAIL
    if email_matches:
        st.session_state["admin_pw"] = st.text_input("관리자 비밀번호", type="password", value=st.session_state["admin_pw"])

is_admin = bool(
    ADMIN_PASSWORD and email_matches and st.session_state["admin_pw"] == ADMIN_PASSWORD
)
with id_c1:
    if is_admin:
        st.success("🔑 관리자 모드 — 확정된 목표를 포함해 모든 항목을 수정할 수 있습니다.")
    elif email_matches and not ADMIN_PASSWORD:
        st.error("서버에 ADMIN_PASSWORD 환경변수가 설정되어 있지 않습니다. 배포 환경변수를 추가해주세요.")
    elif email_matches:
        st.warning("관리자 이메일은 맞지만 비밀번호가 아직 틀렸거나 입력되지 않았습니다.")
    else:
        st.caption("일반 참여자 모드 — 확정 전 목표는 자유롭게 관리하고, 확정 이후 변경은 관리자만 가능합니다.")

st.divider()

# ── 비전 캐스케이드 (기본은 접어둠 — 펼쳐야 보이게) ──────────────
TIERS = [
    {
        "name": "OWM", "sub": "그룹 비전",
        "vision": "입점 브랜드사의 성장 + OWM 각 매장 매출의 비약적 성장.\n\n미션: 이를 위한 전 프로세스의 개발과 효율화.",
        "goal": "입점 브랜드·매장 매출 신장률 **극대화** (구체 목표치 확정 필요)",
    },
    {
        "name": "LSP", "sub": "라이프스타일프로젝트",
        "vision": "옵티마 약국체인 보유. OWM 자회사로서 그룹 비전을 실행 가능한 사업 단위로 전환.",
        "goal": "기업가치 **약 1,000억** → 내년 **2,000~3,000억**",
    },
    {
        "name": "브랜드슬램", "sub": "실행 조직",
        "vision": "미국·중국 SNS 인플루언서 마케팅 + 현지 계정 운영, AX 자동화 파이프라인, 유통·라이브커머스 확장.",
        "goal": "연 **120억** 예산 집행 + 전략 10 / 준전략 50~70 브랜드 유지 + 노스스타 KPI(매장·브랜드 매출 신장률, 콘텐츠 종합평가 성장률) 달성",
    },
]

with st.expander("🗺️ 비전 캐스케이드 (OWM → LSP → 브랜드슬램) — 펼쳐서 보기", expanded=False):
    for t in TIERS:
        with st.container(border=True):
            c1, c2, c3 = st.columns([1, 2.3, 2.3])
            with c1:
                st.markdown(f"**{t['name']}**")
                st.caption(t["sub"])
            with c2:
                st.markdown("**비전 · 미션**")
                st.write(t["vision"])
            with c3:
                st.markdown("**결과적 목표**")
                st.markdown(t["goal"])

    st.write("")
    n1, n2, n3 = st.columns(3)
    n1.metric("연간 OWM 마케팅 예산", "120억")
    n2.metric("월 집행 목표", "10억 / 월")
    n3.metric("전략 / 준전략 브랜드", "10 / 50~70")

st.divider()

# ── 데이터 로드 ───────────────────────────────────────────────
PEOPLE_ORDER = ["양혜준", "김선재", "정다영", "구정회", "박솔 이사", "미국 리드"]
CADENCE_LABEL = {"weekly": "주간", "monthly": "월간", "quarterly": "분기", "once": "1회성"}
CADENCE_KEYS = ["monthly", "weekly", "quarterly", "once"]
BADGE = {
    "ok":    ("#e9f7ee", "#15803d"),
    "watch": ("#fdf2e0", "#b45309"),
    "late":  ("#fdecec", "#dc2626"),
    "blue":  ("#eaf0fd", "#2451c4"),
}

@st.cache_data(ttl=10)
def load_org():
    rows = SUPA.table("okr_org").select("*").execute().data
    return {r["person"]: r for r in rows}

@st.cache_data(ttl=10)
def load_items():
    return SUPA.table("okr_items").select("*").order("created_at").execute().data

def refresh():
    load_org.clear(); load_items.clear(); st.rerun()

# ── 기간/페이스 계산 (데이터 로드보다 먼저 정의 — 롤오버에서 사용) ──
def start_of_week(d): return d - timedelta(days=d.weekday())
def end_of_week(d): return start_of_week(d) + timedelta(days=6)
def start_of_month(d): return d.replace(day=1)
def end_of_month(d):
    nxt = d.replace(day=28) + timedelta(days=4)
    return nxt.replace(day=1) - timedelta(days=1)
def start_of_quarter(d):
    q = (d.month - 1) // 3
    return date(d.year, q * 3 + 1, 1)
def end_of_quarter(d):
    q = (d.month - 1) // 3
    if q == 3:
        return date(d.year, 12, 31)
    return date(d.year, (q + 1) * 3 + 1, 1) - timedelta(days=1)

def period_bounds(cadence, ref):
    if cadence == "weekly": return start_of_week(ref), end_of_week(ref)
    if cadence == "monthly": return start_of_month(ref), end_of_month(ref)
    if cadence == "quarterly": return start_of_quarter(ref), end_of_quarter(ref)
    return ref, ref

def next_period_bounds(cadence, ref):
    """'다음 기간' 기준점을 하나 앞으로 밀어서 그 기간의 시작/끝을 계산한다."""
    if cadence == "weekly":
        ref2 = ref + timedelta(days=7)
    elif cadence == "monthly":
        ref2 = date(ref.year + (1 if ref.month == 12 else 0), (ref.month % 12) + 1, 1)
    elif cadence == "quarterly":
        q = (ref.month - 1) // 3
        ref2 = date(ref.year + (1 if q == 3 else 0), 1 if q == 3 else (q + 1) * 3 + 1, 1)
    else:
        ref2 = ref
    return period_bounds(cadence, ref2)

def period_for(item, ref=None):
    ref = ref or date.today()
    c = item["cadence"]
    if c in ("weekly", "monthly", "quarterly"):
        return period_bounds(c, ref)
    d = item.get("once_date")
    if d:
        d = date.fromisoformat(d) if isinstance(d, str) else d
        return d, d
    return ref, ref

def _parse_d(v):
    if v is None: return None
    return date.fromisoformat(v) if isinstance(v, str) else v

def is_achieved(item):
    target = float(item.get("target_qty") or 0)
    progress = float(item.get("progress") or 0)
    return target > 0 and progress >= target

def close_one_period(it):
    """승인된 기간을 okr_period_log에 스냅샷으로 박제하고 다음 기간으로 리셋한다."""
    today = date.today()
    stored_end = _parse_d(it.get("period_end")) or today
    old_start = _parse_d(it.get("period_start")) or stored_end
    SUPA.table("okr_period_log").insert({
        "person": it["person"], "category": it["category"], "title": it["title"],
        "cadence": it["cadence"], "period_start": old_start.isoformat(), "period_end": stored_end.isoformat(),
        "target_qty": it["target_qty"], "progress": it["progress"],
        "achieved": is_achieved(it),
    }).execute()
    new_start, new_end = period_for(it, today)
    new_target = it.get("next_target_qty") if it.get("next_target_qty") is not None else it["target_qty"]
    new_progress = it["progress"] if it.get("cumulative") else 0
    new_ratio = (new_progress / new_target) if new_target else 0
    SUPA.table("okr_items").update({
        "target_qty": new_target, "progress": new_progress, "next_target_qty": None,
        "period_start": new_start.isoformat(), "period_end": new_end.isoformat(),
        "confirmed": False, "confirmed_at": None, "needs_clarification": False,
        "closing_submitted": False, "closing_report": None, "closing_submitted_at": None,
        "closing_approved": False, "closing_approved_at": None,
        "achieved_at": (today.isoformat() if new_ratio >= 1 else None),
        "updated_at": datetime.utcnow().isoformat(),
    }).eq("id", it["id"]).execute()

def run_period_rollover(items):
    """기간이 끝났어도 '마감 승인(closing_approved)'이 되기 전까지는 절대 자동으로 넘어가지 않는다.
    실제 마감은 담당자 리포트 제출 → 상사(관리자) 승인이 있어야만 발생한다.
    이 함수는 이미 승인된 건만 정리(archive+reset)하고, period_start/end가 비어 있는 신규 항목만 초기화한다."""
    today = date.today()
    rolled = False
    for it in items:
        if it["cadence"] not in ("weekly", "monthly", "quarterly"):
            continue
        stored_end = _parse_d(it.get("period_end"))
        if stored_end is None:
            start, end = period_for(it, today)
            SUPA.table("okr_items").update({"period_start": start.isoformat(), "period_end": end.isoformat()}).eq("id", it["id"]).execute()
            continue
        if stored_end >= today:
            continue  # 아직 현재 기간 — 손대지 않음
        if not it.get("closing_approved"):
            continue  # 기간은 지났지만 아직 마감 승인 전 — '마감 PT 대기' 상태로 그대로 둔다
        close_one_period(it)
        rolled = True
    return rolled

def is_overdue_for_closing(item):
    if item["cadence"] not in ("weekly", "monthly", "quarterly"):
        return False
    end = _parse_d(item.get("period_end"))
    return bool(end) and end < date.today() and not item.get("closing_approved")

ORG = load_org()
_raw_items = load_items()
if run_period_rollover(_raw_items):
    load_items.clear()
ITEMS = load_items()

def pace_status(item, today=None):
    today = today or date.today()
    target = float(item.get("target_qty") or 0)
    progress = float(item.get("progress") or 0)
    if target <= 0:
        return ("blue", "수치 미확정")
    start, end = period_for(item, today)
    total = max(1, (end - start).days + 1)
    elapsed = min(total, max(0, (today - start).days + 1))
    expected_ratio = elapsed / total
    actual_ratio = progress / target
    if actual_ratio >= 1:
        return ("ok", "달성")
    if today > end:
        return ("late", "기한 초과")
    if actual_ratio < expected_ratio * 0.8:
        return ("late", "지연")
    if actual_ratio < expected_ratio * 0.95:
        return ("watch", "주의")
    return ("ok", "정상 진행")

def deadline_label(item, today=None):
    _, end = period_for(item, today or date.today())
    return end.strftime("%m/%d")

def progress_pct(item):
    t = float(item.get("target_qty") or 0)
    if t <= 0: return 0
    return max(0, min(100, round(float(item.get("progress") or 0) / t * 100)))

def badge_html(key, label, extra=""):
    bg, fg = BADGE[key]
    return (f"<span style='background:{bg};color:{fg};padding:2px 9px;border-radius:20px;"
            f"font-size:11px;font-weight:700;white-space:nowrap'>{label}</span>{extra}")

def bar_html(pct, key):
    _, fg = BADGE[key]
    return (f"<div style='background:#eceef1;border-radius:4px;height:6px;width:90px'>"
            f"<div style='background:{fg};width:{pct}%;height:100%;border-radius:4px'></div></div>"
            f"<span style='font-size:11px;color:#888'>{pct}%</span>")

def items_of(person):
    return [it for it in ITEMS if it["person"] == person]

# ── 담당자 선택 ───────────────────────────────────────────────
st.subheader("개인별 OKR · KPI 관리")
sel_col, _ = st.columns([2, 5])
selected = sel_col.selectbox("담당자", ["전체 보기"] + PEOPLE_ORDER, index=1)

# ── 상위 Objective/KR 고정 헤더 ───────────────────────────────
def render_org_card(name):
    org = ORG.get(name, {})
    krs = org.get("krs") or []
    with st.container(border=True):
        pending_tag = "  ·  🕗 채용예정" if org.get("pending") else ""
        st.caption(f"{name} · {org.get('tag','')}{pending_tag}")
        st.markdown(f"**Objective** · {org.get('objective') or '_미입력_'}")
        for i, kr in enumerate(krs, 1):
            warn = "⚠️" in kr
            st.markdown(f"{'⚠️ ' if warn else ''}**KR{i}** · {kr}")
        with st.expander("Objective / KR 수정" + ("" if is_admin else " (관리자 전용)"), expanded=False):
            if not is_admin:
                st.info("Objective/KR 변경은 관리자(jhw@slam-global.com)만 할 수 있어요.")
            else:
                with st.form(f"org_form_{name}"):
                    new_obj = st.text_input("Objective", value=org.get("objective", ""))
                    new_krs_txt = st.text_area("KR (한 줄에 하나씩)", value="\n".join(krs), height=110)
                    new_email = st.text_input("담당자 이메일 (마감 알림 발송용)", value=org.get("email") or "")
                    if st.form_submit_button("저장"):
                        krs_list = [x.strip() for x in new_krs_txt.split("\n") if x.strip()]
                        SUPA.table("okr_org").update(
                            {"objective": new_obj.strip(), "krs": krs_list, "email": new_email.strip(),
                             "updated_at": datetime.utcnow().isoformat()}
                        ).eq("person", name).execute()
                        st.success("저장했습니다."); refresh()

if selected != "전체 보기":
    render_org_card(selected)
else:
    gcols = st.columns(2)
    for i, name in enumerate(PEOPLE_ORDER):
        with gcols[i % 2]:
            org = ORG.get(name, {})
            with st.container(border=True):
                st.caption(f"{name} · {org.get('tag','')}")
                st.markdown(f"**Objective** · {org.get('objective') or '_미입력_'}")
                for j, kr in enumerate(org.get("krs") or [], 1):
                    st.markdown(f"- KR{j} · {kr}")

# ── 요약 카드 ─────────────────────────────────────────────────
scoped = items_of(selected) if selected != "전체 보기" else ITEMS
total = len(scoped)
late_n = sum(1 for it in scoped if pace_status(it)[0] == "late")
watch_n = sum(1 for it in scoped if pace_status(it)[0] == "watch")
confirmed_n = sum(1 for it in scoped if it["confirmed"])
achieved_n = sum(1 for it in scoped if is_achieved(it))
overdue_n = sum(1 for it in scoped if is_overdue_for_closing(it))

m = st.columns(6)
m[0].metric("관리 중 업무", f"{total}건")
m[1].metric("지연", f"{late_n}건", delta="확인 필요" if late_n else None, delta_color="inverse")
m[2].metric("주의", f"{watch_n}건", delta="확인 필요" if watch_n else None, delta_color="inverse")
m[3].metric("확정된 목표", f"{confirmed_n}/{total}")
m[4].metric("🏆 달성", f"{achieved_n}건")
m[5].metric("⏰ 마감 PT 대기", f"{overdue_n}건", delta="승인 필요" if overdue_n else None, delta_color="inverse")

if overdue_n and is_admin:
    st.warning(f"⏰ **마감 PT 승인이 필요한 항목이 {overdue_n}건** 있습니다. '🗓 마감 PT' 탭에서 담당자 리포트를 확인하고 승인해주세요.")
elif overdue_n:
    st.warning(f"⏰ 기간이 끝난 목표 {overdue_n}건이 상사 확인을 기다리고 있어요. '🗓 마감 PT' 탭에서 리포트를 제출해주세요.")

# ── 탭 (탭마다 폼 key가 겹치지 않도록 ctx를 구분해서 넘김) ────
tab_okr, tab_list, tab_cal, tab_late, tab_week, tab_confirmed, tab_achieved, tab_closing, tab_history, tab_upload, tab_gallery, tab_jd = st.tabs(
    ["OKR표", "리스트로 보기", "캘린더 보기", "미달성 KPI 보기", "이번주 수량체크", "협의된 목표 보기", "🏆 달성한 KPI", "🗓 마감 PT", "📜 지난 기록", "📥 리포트 업로드", "🖼 첨부 모아보기", "📄 직무내역서"]
)

def build_payload_with_achievement(it, n_qty, n_progress):
    """진행률이 목표를 넘기면 achieved_at 자동 기록, 다시 내려가면 해제."""
    ratio = (n_progress / n_qty) if n_qty > 0 else 0
    achieved_at = it.get("achieved_at")
    if ratio >= 1 and not achieved_at:
        achieved_at = date.today().isoformat()
    elif ratio < 1 and achieved_at:
        achieved_at = None
    return achieved_at

def item_detail_form(it, ctx):
    """ctx: 'okr' | 'late' | 'cal' | 'achieved' — 같은 항목이 여러 탭에 동시에 나와도 폼 key가 겹치지 않게 함."""
    locked = it["confirmed"] and not is_admin
    label = "▾ 상세 · 수정" + (" 🔒" if locked else "")
    recurring = it["cadence"] in ("weekly", "monthly", "quarterly")
    with st.expander(label):
        if locked:
            st.info("🔒 이 목표는 이미 협의·확정되었습니다. 변경은 관리자(jhw@slam-global.com)만 할 수 있어요.")
            st.write(f"목표 {it['target_qty']} {it['unit']} · 진행 {it['progress']} · 마감 {deadline_label(it)}")
            if it.get("note"):
                st.caption(it["note"])
            atts = it.get("attachments") or []
            if atts:
                st.markdown("**📎 첨부 (보기 전용 — 확정된 목표라 추가/삭제 불가)**")
                for a in atts:
                    if a.get("type") == "photo":
                        st.image(a["url"], width=180, caption=a.get("label"))
                    else:
                        st.markdown(f"🔗 [{a.get('label') or a['url']}]({a['url']})")
            return

        if recurring and it.get("period_start") and it.get("period_end"):
            st.caption(f"📅 이번 기간 · {it['period_start']} ~ {it['period_end']} (지난 기간은 자동으로 '지난 기록'에 박제되어 수정할 수 없어요)")
        if is_overdue_for_closing(it):
            st.warning("⏰ 이 기간은 이미 끝났습니다. 여기서 숫자는 계속 고칠 수 있지만, **'🗓 마감 PT' 탭에서 리포트 제출 → 관리자 승인**이 있어야 실제로 마감됩니다.")
        if it.get("next_target_qty") is not None:
            st.caption(f"⏭ 다음 기간부터 목표 **{it['next_target_qty']} {it['unit']}**(으)로 바뀔 예정입니다.")

        with st.form(f"edit_{ctx}_{it['id']}"):
            e1, e2, e3, e4 = st.columns(4)
            n_title = e1.text_input("업무명", value=it["title"])
            n_unit = e3.text_input("단위", value=it["unit"] or "건")
            n_cadence = e4.selectbox("주기", CADENCE_KEYS,
                                      index=CADENCE_KEYS.index(it["cadence"]),
                                      format_func=lambda x: CADENCE_LABEL[x])

            if recurring:
                e2.caption("목표 수량 ↓ (적용 시점 선택)")
                n_qty = e2.number_input("새 목표 수량", value=float(it["next_target_qty"] if it.get("next_target_qty") is not None else it["target_qty"]), min_value=0.0, label_visibility="collapsed")
                apply_when = e2.radio("적용 시점", ["이번 기간에 즉시 적용", "다음 기간부터 적용"], horizontal=False, label_visibility="collapsed")
            else:
                n_qty = e2.number_input("목표 수량", value=float(it["target_qty"]), min_value=0.0)
                apply_when = "이번 기간에 즉시 적용"

            p1, p2 = st.columns(2)
            n_progress = p1.number_input("현재 진행 (이번 기간)", value=float(it["progress"]), min_value=0.0)
            n_once = None
            n_cumulative = it.get("cumulative", False)
            if n_cadence == "once":
                default_once = date.fromisoformat(it["once_date"]) if it.get("once_date") else date.today()
                n_once = p2.date_input("마감일(1회성)", value=default_once)
            elif recurring and is_admin:
                n_cumulative = p2.checkbox("누적형 KPI (기간이 바뀌어도 진행 수치를 초기화하지 않음)", value=it.get("cumulative", False))
            else:
                if recurring:
                    p2.caption(("누적형 KPI (관리자만 변경 가능)" if not n_cumulative else "🔁 누적형 — 기간이 바뀌어도 리셋되지 않음"))

            n_note = st.text_area("메모 / 상세 설명", value=it.get("note") or "")

            c1, c2 = st.columns(2)
            if is_admin:
                n_confirmed = c1.checkbox("협의된 목표로 확정 (관리자 전용)", value=it["confirmed"])
            else:
                c1.caption("🔒 확정 여부는 관리자만 바꿀 수 있어요.")
                n_confirmed = it["confirmed"]
            n_clarify = c2.checkbox("💬 설명필요 표시", value=it["needs_clarification"])

            b1, b2, b3 = st.columns([1, 1, 3])
            save = b1.form_submit_button("저장", type="primary")
            delete = b2.form_submit_button("삭제")

            if save:
                immediate = (apply_when == "이번 기간에 즉시 적용")
                new_target_now = n_qty if immediate else it["target_qty"]
                new_next_target = None if immediate else n_qty

                achieved_at = build_payload_with_achievement(it, new_target_now, n_progress)
                payload = {
                    "title": n_title.strip(), "unit": n_unit.strip() or "건",
                    "cadence": n_cadence, "progress": n_progress, "note": n_note,
                    "needs_clarification": n_clarify, "achieved_at": achieved_at,
                    "once_date": n_once.isoformat() if n_once else None,
                    "target_qty": new_target_now, "next_target_qty": new_next_target,
                    "cumulative": n_cumulative,
                    "updated_at": datetime.utcnow().isoformat(),
                }
                if n_cadence != it["cadence"] and n_cadence in ("weekly", "monthly", "quarterly"):
                    ns, ne = period_bounds(n_cadence, date.today())
                    payload["period_start"] = ns.isoformat(); payload["period_end"] = ne.isoformat()
                if is_admin:
                    payload["confirmed"] = n_confirmed
                    if n_confirmed and not it["confirmed"]:
                        payload["confirmed_at"] = date.today().isoformat()
                    elif not n_confirmed:
                        payload["confirmed_at"] = None
                SUPA.table("okr_items").update(payload).eq("id", it["id"]).execute()
                if achieved_at and not it.get("achieved_at"):
                    st.balloons()
                    st.success("🏆 목표 달성! 축하합니다.")
                elif not immediate:
                    st.success(f"저장했습니다. 목표 변경은 다음 기간({CADENCE_LABEL[n_cadence]})부터 적용됩니다.")
                else:
                    st.success("저장했습니다.")
                refresh()
            if delete:
                SUPA.table("okr_items").delete().eq("id", it["id"]).execute()
                st.warning("삭제했습니다."); refresh()

        st.markdown("**📎 링크 · 완료 스크린샷**")
        render_attachments_editor(it, key_prefix=f"{ctx}_{it['id']}")

GRID_WITH_PERSON = "0.7fr 3.1fr 0.8fr 1.1fr 1.3fr 0.7fr 1.2fr"
GRID_NO_PERSON = "3.6fr 0.8fr 1.1fr 1.3fr 0.7fr 1.2fr"

def render_header_row(show_person=False):
    cols_def = GRID_WITH_PERSON if show_person else GRID_NO_PERSON
    labels = (["담당자", "업무", "주기", "목표", "진행률", "마감", "상태"] if show_person
              else ["업무", "주기", "목표", "진행률", "마감", "상태"])
    cells = "".join(f"<div>{l}</div>" for l in labels)
    st.markdown(f"<div class='okr-hdr' style='grid-template-columns:{cols_def}'>{cells}</div>", unsafe_allow_html=True)

def render_item_row(it, ctx, show_person=False):
    key, label = pace_status(it)
    pct = progress_pct(it)
    dl = deadline_label(it)
    crown = " 👑" if is_achieved(it) else ""
    extra = " 💬" if it["needs_clarification"] else ""
    qty_str = f"{it['progress']}/{it['target_qty']} {it['unit']}" if it["target_qty"] > 0 else "미확정"
    conf_badge = " " + badge_html("blue", "확정") if it["confirmed"] else ""
    overdue_badge = " " + badge_html("late", "마감PT 대기") if is_overdue_for_closing(it) else ""
    cols_def = GRID_WITH_PERSON if show_person else GRID_NO_PERSON
    person_cell = f"<div class='okr-person'>{it['person']}</div>" if show_person else ""
    html = (
        f"<div class='okr-row-html' style='grid-template-columns:{cols_def}'>"
        f"{person_cell}"
        f"<div><span class='okr-title'>{it['title']}{crown}</span><span class='okr-cat'>{it['category']}{extra}</span></div>"
        f"<div class='okr-num'>{CADENCE_LABEL[it['cadence']]}</div>"
        f"<div class='okr-num'>{qty_str}</div>"
        f"<div>{bar_html(pct, key)}</div>"
        f"<div class='okr-num'>{dl}</div>"
        f"<div>{badge_html(key, label)}{conf_badge}{overdue_badge}</div>"
        f"</div>"
    )
    st.markdown(html, unsafe_allow_html=True)
    item_detail_form(it, ctx)

# ── OKR표 ───────────────────────────────────────────────────
with tab_okr:
    people_shown = [selected] if selected != "전체 보기" else PEOPLE_ORDER
    if selected != "전체 보기":
        with st.expander("+ 새 업무 추가"):
            with st.form(f"add_{selected}", clear_on_submit=True):
                a1, a2, a3, a4 = st.columns(4)
                t_title = a1.text_input("업무명")
                t_qty = a2.number_input("수량", min_value=0.0, value=0.0)
                t_unit = a3.text_input("단위", value="건")
                t_cadence = a4.selectbox("주기", CADENCE_KEYS, format_func=lambda x: CADENCE_LABEL[x])
                st.caption("🔗 링크·📷 사진은 선택사항이에요 (나중에 상세보기에서도 계속 추가할 수 있습니다)")
                l1, l2 = st.columns(2)
                t_link = l1.text_input("링크 URL (선택)")
                t_link_label = l2.text_input("링크 설명 (선택)")
                t_photo = st.file_uploader("사진 첨부 (선택)", type=["png", "jpg", "jpeg", "webp"], key="new_item_photo")
                if st.form_submit_button("추가"):
                    if t_title.strip():
                        attachments = []
                        if t_link.strip():
                            attachments.append({"type": "link", "url": t_link.strip(), "label": t_link_label.strip(), "added_at": datetime.utcnow().isoformat()})
                        if t_photo is not None:
                            try:
                                url = upload_photo(t_photo, selected)
                                attachments.append({"type": "photo", "url": url, "label": t_photo.name, "added_at": datetime.utcnow().isoformat()})
                            except Exception as e:
                                st.error(f"사진 업로드 실패: {e}")
                        payload = {
                            "person": selected, "title": t_title.strip(), "target_qty": t_qty,
                            "unit": t_unit.strip() or "건", "cadence": t_cadence, "attachments": attachments,
                        }
                        if t_cadence in ("weekly", "monthly", "quarterly"):
                            ns, ne = period_bounds(t_cadence, date.today())
                            payload["period_start"] = ns.isoformat(); payload["period_end"] = ne.isoformat()
                        SUPA.table("okr_items").insert(payload).execute()
                        st.success("추가했습니다."); refresh()
                    else:
                        st.error("업무명을 입력해주세요.")
    for p in people_shown:
        p_items = items_of(p)
        if selected == "전체 보기":
            st.markdown(f"#### {p} · {ORG.get(p,{}).get('tag','')}")
        if not p_items:
            st.caption("아직 등록된 업무가 없습니다." + ("" if selected != "전체 보기" else " (위에서 담당자를 선택해 추가하세요)"))
            st.write("")
            continue
        render_header_row(show_person=False)
        for it in p_items:
            render_item_row(it, ctx="okr", show_person=False)

# ── 리스트로 보기 ───────────────────────────────────────────
with tab_list:
    all_items = items_of(selected) if selected != "전체 보기" else ITEMS
    if not all_items:
        st.caption("표시할 항목이 없습니다.")
    else:
        all_items_sorted = sorted(all_items, key=lambda x: (pace_status(x)[0] != "late", pace_status(x)[0] != "watch"))
        rows = []
        for it in all_items_sorted:
            key, label = pace_status(it)
            rows.append({
                "담당자": it["person"], "업무": ("👑 " if is_achieved(it) else "") + it["title"], "카테고리": it["category"],
                "주기": CADENCE_LABEL[it["cadence"]],
                "목표": f"{it['target_qty']} {it['unit']}" if it["target_qty"] > 0 else "미확정",
                "진행": f"{it['progress']} ({progress_pct(it)}%)",
                "마감": deadline_label(it), "상태": label,
                "확정": "✅" if it["confirmed"] else "",
            })
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

# ── 캘린더 보기 ─────────────────────────────────────────────
with tab_cal:
    cal_items = items_of(selected) if selected != "전체 보기" else ITEMS
    today = date.today()
    month_offset = st.session_state.get("okr_cal_offset", 0)
    cnav = st.columns([1, 3, 1])
    if cnav[0].button("‹ 이전달"):
        st.session_state["okr_cal_offset"] = month_offset - 1; st.rerun()
    if cnav[2].button("다음달 ›"):
        st.session_state["okr_cal_offset"] = month_offset + 1; st.rerun()

    base_month = today.month - 1 + month_offset
    base_year = today.year + base_month // 12
    base_month = base_month % 12 + 1
    base = date(base_year, base_month, 1)
    cnav[1].markdown(f"<div style='text-align:center;font-weight:800;font-size:16px'>{base.year}년 {base.month}월</div>", unsafe_allow_html=True)

    grid_start = start_of_week(base)
    grid_end = end_of_week(end_of_month(base))

    day_map = {}       # 마감일 -> 항목들
    crown_map = {}     # 달성일(achieved_at) -> 항목들
    for it in cal_items:
        c = it["cadence"]
        if c == "once":
            d = it.get("once_date")
            if d:
                d = date.fromisoformat(d) if isinstance(d, str) else d
                if grid_start <= d <= grid_end:
                    day_map.setdefault(d, []).append(it)
        else:
            cursor = grid_start
            seen = set()
            while cursor <= grid_end:
                _, end = period_for(it, cursor)
                if grid_start <= end <= grid_end and end not in seen:
                    day_map.setdefault(end, []).append(it)
                    seen.add(end)
                cursor += timedelta(days=7 if c == "weekly" else 30)
        if it.get("achieved_at"):
            ad = date.fromisoformat(it["achieved_at"]) if isinstance(it["achieved_at"], str) else it["achieved_at"]
            if grid_start <= ad <= grid_end:
                crown_map.setdefault(ad, []).append(it)

    dows = ["월", "화", "수", "목", "금", "토", "일"]
    html = "<div style='display:grid;grid-template-columns:repeat(7,1fr);gap:6px'>"
    for dw in dows:
        html += f"<div style='font-size:11px;color:#9ca3af;text-align:center;font-weight:700'>{dw}</div>"
    cursor = grid_start
    while cursor <= grid_end:
        out = cursor.month != base.month
        is_today = cursor == today
        has_crown = cursor in crown_map
        bg = "#fffbea" if has_crown else ("#f7f7f8" if out else "#fff")
        border = "1.5px solid #d4a017" if has_crown else ("1px solid #16181d" if is_today else "1px solid #e7e8ec")
        chips = ""
        if has_crown:
            chips += "<div style='font-size:14px'>👑</div>"
        for it in day_map.get(cursor, [])[:3]:
            k, _ = pace_status(it)
            bg2, fg2 = BADGE[k]
            chips += (f"<div style='background:{bg2};color:{fg2};font-size:10px;padding:2px 5px;"
                      f"border-radius:4px;margin-top:2px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis'>"
                      f"{it['person']} · {it['title']}</div>")
        extra = len(day_map.get(cursor, [])) - 3
        if extra > 0:
            chips += f"<div style='font-size:10px;color:#9ca3af'>+{extra}건 더</div>"
        html += (f"<div style='background:{bg};border:{border};border-radius:8px;padding:6px;min-height:74px'>"
                 f"<div style='font-size:11px;font-weight:700;color:#6b7280'>{cursor.day}</div>{chips}</div>")
        cursor += timedelta(days=1)
    html += "</div>"
    st.markdown(html, unsafe_allow_html=True)
    st.caption("👑 표시된 날짜 = 그날 목표를 달성한 기록이 있어요.")

    st.write("")
    pick = st.date_input("특정 날짜의 마감/달성 항목 보기", value=today)
    day_items = day_map.get(pick, [])
    crown_items = crown_map.get(pick, [])
    if crown_items:
        st.markdown("**🏆 이 날짜에 달성한 목표**")
        for it in crown_items:
            render_item_row(it, ctx="cal_crown", show_person=(selected == "전체 보기"))
    if day_items:
        st.markdown(f"**{pick.strftime('%Y-%m-%d')} 마감 항목**")
        for it in day_items:
            render_item_row(it, ctx="cal", show_person=(selected == "전체 보기"))
    if not day_items and not crown_items:
        st.caption("해당 날짜에 마감·달성 기록이 없습니다.")

# ── 미달성 KPI 보기 ─────────────────────────────────────────
with tab_late:
    scope_items = items_of(selected) if selected != "전체 보기" else ITEMS
    risky = [it for it in scope_items if pace_status(it)[0] in ("late", "watch")]
    if not risky:
        st.success("🎉 현재 지연·주의 항목이 없습니다.")
    else:
        risky.sort(key=lambda x: pace_status(x)[0] != "late")
        for it in risky:
            start, end = period_for(it)
            today = date.today()
            total = max(1, (end - start).days + 1)
            elapsed = min(total, max(0, (today - start).days + 1))
            expected = round(float(it["target_qty"]) * elapsed / total) if it["target_qty"] else 0
            key, label = pace_status(it)
            with st.container(border=True):
                cc = st.columns([3, 2, 1.4])
                cc[0].markdown(f"**{it['person']} · {it['title']}**  \n<span style='font-size:11px;color:#9ca3af'>{it['category']}</span>", unsafe_allow_html=True)
                cc[1].write(f"실제 {it['progress']} / 기대치 ≈ {expected} {it['unit']} (목표 {it['target_qty']})")
                cc[2].markdown(badge_html(key, label), unsafe_allow_html=True)
                item_detail_form(it, ctx="late")

# ── 이번주 수량체크 ─────────────────────────────────────────
with tab_week:
    scope_items = items_of(selected) if selected != "전체 보기" else ITEMS
    weekly_items = [it for it in scope_items if it["cadence"] == "weekly"]
    if not weekly_items:
        st.caption("주간 단위로 관리 중인 업무가 없습니다. OKR표에서 주기를 '주간'으로 등록해보세요.")
    else:
        wk_start = start_of_week(date.today())
        st.caption(f"이번 주 · {wk_start.strftime('%m/%d')} 시작")
        for it in weekly_items:
            key, label = pace_status(it)
            with st.form(f"week_{it['id']}"):
                wc = st.columns([3, 1.2, 1.4, 1])
                wc[0].markdown(f"**{it['person']} · {it['title']}**  \n<span style='font-size:11px;color:#9ca3af'>이번 주 목표 {it['target_qty']} {it['unit']}</span>", unsafe_allow_html=True)
                new_val = wc[1].number_input("진행", value=float(it["progress"]), min_value=0.0, label_visibility="collapsed")
                wc[2].markdown(bar_html(progress_pct(it), key) + " " + badge_html(key, label), unsafe_allow_html=True)
                if wc[3].form_submit_button("저장"):
                    achieved_at = build_payload_with_achievement(it, it["target_qty"], new_val)
                    SUPA.table("okr_items").update(
                        {"progress": new_val, "achieved_at": achieved_at, "updated_at": datetime.utcnow().isoformat()}
                    ).eq("id", it["id"]).execute()
                    st.success("저장했습니다."); refresh()

# ── 협의된 목표 보기 ─────────────────────────────────────────
with tab_confirmed:
    scope_items = items_of(selected) if selected != "전체 보기" else ITEMS
    confirmed_items = [it for it in scope_items if it["confirmed"]]
    if not confirmed_items:
        st.info("아직 '확정'된 목표가 없습니다. 상세보기에서 **협의된 목표로 확정**(관리자 전용)을 켜면 여기 표시됩니다.")
    else:
        rows = [{
            "담당자": it["person"], "업무": it["title"], "주기": CADENCE_LABEL[it["cadence"]],
            "목표": f"{it['target_qty']} {it['unit']}", "마감": deadline_label(it),
            "확정일": it.get("confirmed_at") or "-",
        } for it in confirmed_items]
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)
        if not is_admin:
            st.caption("🔒 확정된 목표의 변경은 관리자(jhw@slam-global.com)만 할 수 있어요.")

# ── 🏆 달성한 KPI ───────────────────────────────────────────
with tab_achieved:
    scope_items = items_of(selected) if selected != "전체 보기" else ITEMS
    achieved_items = [it for it in scope_items if is_achieved(it)]
    achieved_items.sort(key=lambda x: x.get("achieved_at") or "", reverse=True)
    if not achieved_items:
        st.caption("아직 목표를 달성한 항목이 없습니다. 진행률이 목표치에 도달하면 여기 자동으로 올라오고, 캘린더에도 👑 표시가 붙어요.")
    else:
        st.markdown(
            "<div style='background:linear-gradient(135deg,#fffbea,#fef3c7);border:1px solid #f2e0b8;"
            "border-radius:12px;padding:14px 18px;margin-bottom:14px;font-weight:700;color:#92400e'>"
            f"🏆 지금까지 {len(achieved_items)}개의 목표를 달성했습니다 — 명예의 전당</div>",
            unsafe_allow_html=True
        )
        for it in achieved_items:
            with st.container(border=True):
                ac = st.columns([3, 2, 1.4])
                ac[0].markdown(f"👑 **{it['person']} · {it['title']}**  \n<span style='font-size:11px;color:#9ca3af'>{it['category']}</span>", unsafe_allow_html=True)
                ac[1].write(f"{it['target_qty']} {it['unit']} 달성 · {CADENCE_LABEL[it['cadence']]}")
                ac[2].markdown(f"<span style='background:#fef3c7;color:#92400e;padding:3px 10px;border-radius:20px;font-size:11px;font-weight:800'>🏆 {it.get('achieved_at') or ''}</span>", unsafe_allow_html=True)
                item_detail_form(it, ctx="achieved")

# ── 🗓 마감 PT (담당자 리포트 제출 → 관리자 승인 → 그때 비로소 마감) ──
with tab_closing:
    scope_items = items_of(selected) if selected != "전체 보기" else ITEMS
    overdue_items = [it for it in scope_items if is_overdue_for_closing(it)]
    st.caption("기간이 끝난 목표는 여기서 리포트를 제출하고, 관리자가 실제 성과를 확인·승인해야 비로소 마감됩니다. "
               "승인 전까지는 숫자를 계속 고칠 수 있어요 (면담하면서 같이 확인하고 조정 가능).")

    if is_admin:
        with st.expander("📧 리마인더 이메일 — 테스트 / 지금 즉시 발송 (관리자 전용)"):
            st.caption("Resend 도메인 인증이 끝났으면 여기서 바로 테스트하거나, 목요일을 기다리지 않고 지금 리마인더를 보낼 수 있어요.")
            tcol1, tcol2 = st.columns([3, 1])
            test_addr = tcol1.text_input("테스트로 받을 이메일", value=ADMIN_EMAIL)
            if tcol2.button("테스트 발송"):
                ok, err = send_via_resend(
                    test_addr, "[브랜드슬램 OKR] 테스트 메일",
                    "이 메일이 도착했다면 Resend 연동이 정상 작동 중입니다.",
                    purpose="test",
                )
                if ok:
                    st.success(f"✅ {test_addr} 로 발송 완료! 받은편지함(스팸함도) 확인해보세요.")
                else:
                    st.error(f"발송 실패: {err}")

            st.divider()
            if st.button("🔔 지금 '이번주 마감 PT' 리마인더 즉시 발송 (전체 대상)"):
                weekly_items_all = [it for it in ITEMS if it["cadence"] == "weekly"]
                by_person_w = {}
                for it in weekly_items_all:
                    by_person_w.setdefault(it["person"], []).append(it)
                sent, skipped = [], []
                for name in PEOPLE_ORDER:
                    org = ORG.get(name, {})
                    email = (org.get("email") or "").strip()
                    if org.get("pending") or not email:
                        skipped.append(f"{name} (이메일 없음/공석)"); continue
                    not_sub = [it for it in by_person_w.get(name, []) if not it.get("closing_submitted")]
                    if not by_person_w.get(name):
                        skipped.append(f"{name} (주간 KPI 없음)"); continue
                    if not not_sub:
                        skipped.append(f"{name} (이미 전부 제출함)"); continue
                    lines = "\n".join(f"  - {it['title']} (진행 {it['progress']}/{it['target_qty']} {it['unit']})" for it in not_sub)
                    body = (f"{name}님, 안녕하세요.\n\n이번 주 마감 PT를 아직 제출하지 않은 항목이 있어요:\n\n{lines}\n\n"
                            "OKR 목표관리 페이지의 '🗓 마감 PT' 탭에서 이번 주 리포트를 작성해 제출해주세요.\n"
                            "- 브랜드슬램 OKR 시스템 (관리자 수동 발송)")
                    ok, err = send_via_resend(email, "[브랜드슬램 OKR] 이번 주 마감 PT 제출 요청", body, purpose="manual_weekly_reminder")
                    (sent if ok else skipped).append(f"{name}{'' if ok else f' (실패: {err})'}")
                st.success(f"발송 완료: {', '.join(sent) if sent else '없음'}")
                if skipped:
                    st.caption("건너뜀/실패: " + ", ".join(skipped))

    if not overdue_items:
        st.success("🎉 마감 PT 대기 중인 항목이 없습니다.")
    else:
        for it in overdue_items:
            with st.container(border=True):
                hc = st.columns([3, 1.6, 1.4])
                hc[0].markdown(f"**{it['person']} · {it['title']}**  \n<span style='font-size:11px;color:#9ca3af'>{it['category']} · 기간 {it.get('period_start')} ~ {it.get('period_end')}</span>", unsafe_allow_html=True)
                hc[1].write(f"실적 {it['progress']} / 목표 {it['target_qty']} {it['unit']}")
                if it["closing_submitted"]:
                    hc[2].markdown(badge_html("watch", "제출됨 · 승인 대기"), unsafe_allow_html=True)
                else:
                    hc[2].markdown(badge_html("late", "리포트 미제출"), unsafe_allow_html=True)

                with st.form(f"closing_{it['id']}"):
                    n_progress = st.number_input("최종 실적 (필요하면 면담하면서 조정)", value=float(it["progress"]), min_value=0.0)
                    n_report = st.text_area("이번 기간 리포트 (무엇을 했는지 간단히)", value=it.get("closing_report") or "",
                                             placeholder="예: 목표 50건 중 42건 완료. 나머지는 다음주로 이월...")
                    b1, b2 = st.columns(2)
                    submit_report = b1.form_submit_button("📝 마감 요청 제출 (상사 확인 요청)")
                    approve = b2.form_submit_button("✅ 마감 승인 (관리자 전용)", type="primary", disabled=not is_admin)

                    if submit_report:
                        SUPA.table("okr_items").update({
                            "progress": n_progress, "closing_report": n_report,
                            "closing_submitted": True, "closing_submitted_at": datetime.utcnow().isoformat(),
                            "updated_at": datetime.utcnow().isoformat(),
                        }).eq("id", it["id"]).execute()
                        st.success("제출했습니다. 관리자 승인을 기다려주세요."); refresh()

                    if approve and is_admin:
                        # 승인 직전 최종 실적을 반영한 뒤, 그 스냅샷으로 즉시 마감 처리
                        SUPA.table("okr_items").update({
                            "progress": n_progress, "closing_report": n_report,
                            "closing_approved": True, "closing_approved_at": datetime.utcnow().isoformat(),
                            "updated_at": datetime.utcnow().isoformat(),
                        }).eq("id", it["id"]).execute()
                        it_for_close = dict(it); it_for_close["progress"] = n_progress
                        close_one_period(it_for_close)
                        st.success("✅ 마감 승인 완료 — 지난 기록으로 이동하고 다음 기간이 시작됩니다.")
                        st.balloons()
                        refresh()


with tab_history:
    q = SUPA.table("okr_period_log").select("*").order("period_end", desc=True).limit(300)
    if selected != "전체 보기":
        q = q.eq("person", selected)
    hist_rows = q.execute().data
    if not hist_rows:
        st.caption("아직 마감된 기간이 없습니다. 주간/월간/분기 목표의 기간이 지나면 여기에 자동으로 기록이 쌓입니다 (수정 불가 — 순수 기록용).")
    else:
        st.caption("지난 기간의 목표·실적 스냅샷입니다. 이미 지나간 기록이라 여기서는 수정할 수 없어요.")
        rows = [{
            "담당자": r["person"], "업무": ("👑 " if r["achieved"] else "") + r["title"],
            "주기": CADENCE_LABEL.get(r["cadence"], r["cadence"]),
            "기간": f"{r['period_start']} ~ {r['period_end']}",
            "목표": f"{r['target_qty']}", "실적": f"{r['progress']}",
            "달성": "✅" if r["achieved"] else "",
        } for r in hist_rows]
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

# ── 📥 리포트 업로드 (문서 → Claude가 KPI 후보 추출 → 검토 → 반영) ──
with tab_upload:
    st.caption(
        "직원이 보내온 업무계획서·KPI 리포트 파일(.docx / .pdf / .txt)을 그대로 올리면, "
        "Claude가 읽고 수량화 가능한 목표를 뽑아 후보로 보여줍니다. "
        "**바로 반영되지 않고, 검토·수정 후 '반영' 버튼을 눌러야만** OKR표에 들어가요. "
        "올린 문서와 추출 결과는 전부 저장되고, 다음 분석 때 같은 담당자의 과거 사례를 참고자료로 함께 사용합니다."
    )
    up_person = st.selectbox("이 리포트는 누구 것인가요?", PEOPLE_ORDER, key="upload_person")
    up_period_choice = st.radio(
        "이 리포트는 어느 기간의 목표인가요?",
        ["이번 기간 (지금 진행 중인 주/달/분기)", "다음 기간 (아직 시작 전)"],
        horizontal=True, key="upload_period_choice",
        help="예: 8월 말에 9월 계획서를 올리는 거라면 '다음 기간'을 선택하세요. 잘못 고르면 엉뚱한 기간에 목표가 꽂힙니다.",
    )
    up_file = st.file_uploader("리포트 파일 업로드", type=["docx", "pdf", "txt"], key="upload_file")
    force_reanalyze = st.checkbox("캐시 무시하고 다시 분석 (같은 파일로 재테스트할 때 체크)", value=False, key="upload_force")

    if up_file is not None and st.button("🧠 Claude로 KPI 추출하기"):
        with st.spinner("문서를 읽고 KPI 후보를 뽑는 중..."):
            raw_text = extract_text_from_upload(up_file)
            h = content_hash_of(raw_text)
            cached = None if force_reanalyze else find_cached_extraction(h)
            if cached is not None:
                items, err = cached, None
                st.info("🔁 예전에 완전히 똑같은 문서를 이미 분석한 적이 있어서, API를 다시 부르지 않고 저장된 결과를 그대로 불러왔습니다 (토큰 절약).")
            elif len(filter_numeric_lines(raw_text)) < 30:
                items, err = [], None
                st.warning("이 문서에서 숫자/수량 관련 문장을 거의 찾지 못했어요. API 호출 없이 건너뛰었습니다 — 수치가 명시된 리포트를 올려주세요.")
            else:
                items, err = call_claude_extract(up_person, raw_text, SUPA.table("okr_report_uploads").select("raw_text,extracted_json")
                                                  .eq("person", up_person).eq("applied", True)
                                                  .order("created_at", desc=True).limit(2).execute().data)
            upload_row = SUPA.table("okr_report_uploads").insert({
                "person": up_person, "filename": up_file.name, "raw_text": raw_text,
                "extracted_json": items, "content_hash": h,
                "admin_email": st.session_state.get("user_email", ""),
            }).execute().data[0]
            if err:
                st.error(f"추출 실패: {err}")
            elif items:
                st.session_state["pending_extract"] = {
                    "upload_id": upload_row["id"], "person": up_person, "items": items,
                    "period_choice": up_period_choice,
                }
                st.success(f"{len(items)}개 KPI 후보를 찾았습니다. 아래에서 검토 후 반영해주세요.")

    pending = st.session_state.get("pending_extract")
    if pending:
        st.markdown(f"### 🔎 추출 결과 검토 · {pending['person']}")
        st.caption(f"적용 시점: **{pending.get('period_choice', '이번 기간')}**")
        existing_lookup = {it["title"].strip().lower(): it for it in items_of(pending["person"])}
        keep_flags = []
        edited_items = []
        for i, item in enumerate(pending["items"]):
            with st.container(border=True):
                k1, k2 = st.columns([0.4, 5.6])
                keep = k1.checkbox("반영", value=True, key=f"keep_{i}")
                with k2:
                    c1, c2, c3, c4 = st.columns(4)
                    e_title = c1.text_input("업무명", value=item.get("title", ""), key=f"et_{i}")
                    e_qty = c2.number_input("목표 수량", value=float(item.get("target_qty") or 0), key=f"eq_{i}")
                    e_unit = c3.text_input("단위", value=item.get("unit", "건"), key=f"eu_{i}")
                    cad_options = ["monthly", "weekly", "quarterly", "once"]
                    default_cad = item.get("cadence") if item.get("cadence") in cad_options else "monthly"
                    e_cad = c4.selectbox("주기", cad_options, index=cad_options.index(default_cad),
                                         format_func=lambda x: CADENCE_LABEL[x], key=f"ec_{i}")
                    e_cat = st.text_input("카테고리", value=item.get("category", "미분류"), key=f"ecat_{i}")
                    if item.get("note"):
                        st.caption(f"📎 근거: {item['note']}")

                    dup = existing_lookup.get(e_title.strip().lower())
                    dup_id = None
                    if dup:
                        if dup["confirmed"] and not is_admin:
                            st.warning(f"⚠️ 동일한 이름의 항목이 이미 있고 **확정되어 잠겨 있습니다** (진행 {dup['progress']}/{dup['target_qty']}). "
                                       "업데이트하려면 관리자로 로그인하거나, 업무명을 다르게 바꿔 새 항목으로 추가하세요.")
                        else:
                            update_mode = st.checkbox(
                                f"⚠️ 이미 같은 이름의 항목이 있어요 (진행 {dup['progress']}/{dup['target_qty']} {dup['unit']}) — "
                                "새로 만들지 않고 그 항목을 업데이트", value=True, key=f"upd_{i}",
                            )
                            dup_id = dup["id"] if update_mode else None
            keep_flags.append(keep)
            edited_items.append({"category": e_cat, "title": e_title, "target_qty": e_qty, "unit": e_unit, "cadence": e_cad, "_dup_id": dup_id})

        bcol1, bcol2 = st.columns([1, 1])
        if bcol1.button("✅ 체크된 항목을 OKR표에 반영", type="primary"):
            inserted, updated = 0, 0
            final_kept_items = []
            use_next_period = str(pending.get("period_choice", "")).startswith("다음")
            for keep, it in zip(keep_flags, edited_items):
                if not keep or not it["title"].strip():
                    continue
                dup_id = it.pop("_dup_id", None)
                if dup_id:
                    SUPA.table("okr_items").update({
                        "category": it["category"], "unit": it["unit"], "cadence": it["cadence"],
                        "target_qty": it["target_qty"], "updated_at": datetime.utcnow().isoformat(),
                    }).eq("id", dup_id).execute()
                    updated += 1
                else:
                    payload = {"person": pending["person"], **it}
                    if it["cadence"] in ("weekly", "monthly", "quarterly"):
                        ns, ne = (next_period_bounds if use_next_period else period_bounds)(it["cadence"], date.today())
                        payload["period_start"] = ns.isoformat(); payload["period_end"] = ne.isoformat()
                    SUPA.table("okr_items").insert(payload).execute()
                    inserted += 1
                final_kept_items.append(it)
            # 실제로 반영된(=사람이 최종 확정한) 항목으로 덮어써야, 다음번 참고자료가 'AI 원본 추측'이 아니라
            # '사람이 고친 정답'이 된다 — 이게 있어야 같은 실수가 반복되지 않는다.
            SUPA.table("okr_report_uploads").update({
                "applied": True, "extracted_json": final_kept_items,
            }).eq("id", pending["upload_id"]).execute()
            st.session_state.pop("pending_extract", None)
            st.success(f"새로 추가 {inserted}건, 기존 항목 업데이트 {updated}건 — {pending['person']}님의 OKR표에 반영했습니다. "
                       "(마감돼서 '지난 기록'으로 넘어간 과거 데이터는 이 작업으로 절대 바뀌지 않습니다)")
            refresh()
        if bcol2.button("취소"):
            st.session_state.pop("pending_extract", None)
            st.rerun()

    st.divider()
    st.caption("📚 지난 업로드 이력")
    hist_up = SUPA.table("okr_report_uploads").select("person,filename,applied,created_at") \
        .order("created_at", desc=True).limit(30).execute().data
    if hist_up:
        st.dataframe(pd.DataFrame([{
            "담당자": r["person"], "파일명": r["filename"],
            "반영 여부": "✅ 반영됨" if r["applied"] else "검토만 함",
            "업로드 시각": r["created_at"],
        } for r in hist_up]), use_container_width=True, hide_index=True)
    else:
        st.caption("아직 업로드 이력이 없습니다.")

# ── 🖼 첨부 모아보기 (모든 링크·사진을 한 곳에서) ──────────────
with tab_gallery:
    all_scope = items_of(selected) if selected != "전체 보기" else ITEMS
    gallery_person = st.selectbox("담당자 필터", ["전체"] + PEOPLE_ORDER,
                                   index=(0 if selected == "전체 보기" else PEOPLE_ORDER.index(selected) + 1))
    flat = []
    for it in ITEMS:
        if gallery_person != "전체" and it["person"] != gallery_person:
            continue
        for a in (it.get("attachments") or []):
            flat.append({**a, "person": it["person"], "item_title": it["title"]})
    flat.sort(key=lambda x: x.get("added_at") or "", reverse=True)

    if not flat:
        st.caption("아직 첨부된 링크·사진이 없습니다. 각 업무의 '상세 · 수정'에서 추가할 수 있어요.")
    else:
        gcol1, gcol2 = st.columns(2)
        photos = [f for f in flat if f.get("type") == "photo"]
        links = [f for f in flat if f.get("type") != "photo"]
        gcol1.metric("📷 사진", f"{len(photos)}건")
        gcol2.metric("🔗 링크", f"{len(links)}건")
        st.divider()

        if photos:
            st.markdown("#### 📷 완료 스크린샷")
            pcols = st.columns(4)
            for i, p in enumerate(photos):
                with pcols[i % 4]:
                    st.image(p["url"], use_container_width=True)
                    st.caption(f"{p['person']} · {p['item_title']}")

        if links:
            st.markdown("#### 🔗 링크")
            for l in links:
                st.markdown(f"- [{l.get('label') or l['url']}]({l['url']}) — *{l['person']} · {l['item_title']}*")

# ── 📄 직무내역서 (채용 온보딩용 PDF 출력) ──────────────────────
with tab_jd:
    st.caption(
        "채용 시 지원자에게 그대로 전달할 수 있는 직무내역서를 PDF로 만듭니다. "
        "카테고리·업무명·수행주기만 담아 심플하게 구성되고, 목표 수량·진행률·급여 등은 포함되지 않습니다."
    )
    jd_person = st.selectbox(
        "어떤 포지션(담당자)의 직무내역서를 만들까요?",
        PEOPLE_ORDER,
        index=(PEOPLE_ORDER.index(selected) if selected in PEOPLE_ORDER else 0),
        key="jd_person_select",
    )
    jd_org = ORG.get(jd_person, {})
    jd_items = items_of(jd_person)

    if not jd_items:
        st.info(f"{jd_person}님에게 등록된 업무가 아직 없습니다. 'OKR표' 탭에서 업무를 먼저 추가해주세요.")
    else:
        st.markdown(f"**{jd_person}** · {jd_org.get('tag','')}")
        if jd_org.get("objective"):
            st.caption(f"직무개요 미리보기 · {jd_org['objective']}")

        st.write("")
        st.markdown("**부분출력용 — 포함할 업무를 개별 체크하세요** (기본: 전체 선택)")

        by_cat_for_ui = {}
        for it in jd_items:
            by_cat_for_ui.setdefault(it["category"] or "미분류", []).append(it)

        jd_check_state = {}
        for cat, tasks in by_cat_for_ui.items():
            st.markdown(f"##### {cat}")
            for it in tasks:
                jd_check_state[it["id"]] = st.checkbox(
                    f"{it['title']} · {CADENCE_LABEL[it['cadence']]}",
                    value=True, key=f"jd_chk_{it['id']}",
                )

        def _jd_group(items_subset):
            g = {}
            for it in items_subset:
                g.setdefault(it["category"] or "미분류", []).append((it["title"], CADENCE_LABEL[it["cadence"]]))
            return g

        st.divider()
        jb1, jb2 = st.columns(2)

        pdf_all = build_jd_pdf(jd_person, jd_org.get("tag", ""), jd_org.get("objective", ""), _jd_group(jd_items))
        jb1.download_button(
            "📄 직무내역서 전체 출력 (PDF)",
            data=pdf_all,
            file_name=f"브랜드슬램_직무내역서_{jd_person}_전체.pdf",
            mime="application/pdf",
            use_container_width=True,
        )

        jd_selected_items = [it for it in jd_items if jd_check_state.get(it["id"])]
        if jd_selected_items:
            pdf_partial = build_jd_pdf(jd_person, jd_org.get("tag", ""), jd_org.get("objective", ""), _jd_group(jd_selected_items))
            jb2.download_button(
                f"📄 선택한 업무만 출력 ({len(jd_selected_items)}건, PDF)",
                data=pdf_partial,
                file_name=f"브랜드슬램_직무내역서_{jd_person}_부분.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        else:
            jb2.caption("체크된 업무가 없어 부분출력을 만들 수 없습니다.")

st.divider()
st.caption("브랜드슬램 내부 참고용 · Supabase(okr_org / okr_items) 실시간 연동 · 확정 목표는 관리자만 수정 가능")
